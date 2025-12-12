#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
简历解析脚本
支持 PDF 和 Word (.docx) 格式简历，通过 LiteLLM 调用 DeepSeek 进行结构化信息提取
"""

import json
import os
import time
from pathlib import Path
from dotenv import load_dotenv
import fitz  # PyMuPDF
from docx import Document  # python-docx
from litellm import completion

# 加载环境变量
env_path = Path(__file__).parent / ".env"
load_dotenv(env_path)

# API 配置
DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY")
DEEPSEEK_BASE_URL = os.getenv("DEEPSEEK_BASE_URL")
AI_MODEL = os.getenv("AI_MODEL", "deepseek/deepseek-chat")
AI_TIMEOUT = int(os.getenv("AI_TIMEOUT", 150))

# 设置 API Key
os.environ["DEEPSEEK_API_KEY"] = DEEPSEEK_API_KEY


def extract_text_from_pdf(pdf_path: str) -> str:
    """
    从 PDF 文件中提取文本
    
    Args:
        pdf_path: PDF 文件路径
        
    Returns:
        提取的文本内容
    """
    doc = fitz.open(pdf_path)
    text_content = []
    
    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text()
        text_content.append(f"--- 第 {page_num + 1} 页 ---\n{text}")
    
    doc.close()
    return "\n\n".join(text_content)


def extract_text_from_docx(docx_path: str) -> str:
    """
    从 Word (.docx) 文件中提取文本
    
    Args:
        docx_path: Word 文件路径
        
    Returns:
        提取的文本内容
    """
    doc = Document(docx_path)
    text_content = []
    
    # 提取段落文本
    for para in doc.paragraphs:
        if para.text.strip():
            text_content.append(para.text)
    
    # 提取表格文本
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                text_content.append(" | ".join(row_text))
    
    return "\n".join(text_content)


def extract_text_from_resume(file_path: str) -> str:
    """
    根据文件类型自动选择提取方法
    
    Args:
        file_path: 简历文件路径（支持 .pdf, .docx）
        
    Returns:
        提取的文本内容
        
    Raises:
        ValueError: 不支持的文件格式
    """
    file_path = Path(file_path)
    suffix = file_path.suffix.lower()
    
    if suffix == '.pdf':
        return extract_text_from_pdf(str(file_path))
    elif suffix == '.docx':
        return extract_text_from_docx(str(file_path))
    elif suffix == '.doc':
        raise ValueError("不支持 .doc 格式，请将文件另存为 .docx 格式")
    else:
        raise ValueError(f"不支持的文件格式: {suffix}，支持的格式: .pdf, .docx")


def parse_resume_with_llm(resume_text: str) -> dict:
    """
    使用 LLM 解析简历文本，提取结构化信息
    
    Args:
        resume_text: 简历文本内容
        
    Returns:
        解析后的结构化信息
    """
    
    system_prompt = """你是一个专业的简历解析助手。你需要从简历文本中提取结构化信息。
请严格按照要求的 JSON 格式返回结果，如果某项信息无法从简历中获取，则填写 null。
日期格式统一为 YYYY-MM-DD。
注意：
1. 教育经历按学历从高到低排序（博士 > 硕士 > 本科）
2. 工作经历和项目经历需要按时间由近及远排序
3. 仅提取 IT 相关的工作和项目经历
4. 确保输出是有效的 JSON 格式
5. 如果某项信息在简历中部分存在，尽可能提取已有信息"""

    user_prompt = f"""请解析以下简历文本，提取结构化信息并返回 JSON 格式结果。

简历文本：
{resume_text}

请返回以下 JSON 格式的结果（所有字段都必须存在，如果信息不存在则填 null）：

{{
    "basic_info": {{
        "name": "姓名",
        "supplier": "所属供应商"
    }},
    "personal_info": {{
        "id_card_number": "身份证号",
        "birth_date": "出生日期（YYYY-MM-DD）",
        "phone": "电话",
        "first_work_date": "首次参加工作时间（YYYY-MM-DD）",
        "first_it_work_date": "首次参加IT领域工作时间（YYYY-MM-DD）",
        "highest_education": "最高学历（本科/硕士/博士）",
        "contract_level": "人员输送当时符合的合同定级"
    }},
    "education": [
        {{
            "degree_type": "学历类型（本科/硕士/博士）",
            "enrollment_date": "入学时间（YYYY-MM-DD）",
            "university": "毕业院校",
            "graduation_date": "毕业时间（YYYY-MM-DD）",
            "major": "专业",
            "diploma_number": "毕业证编号",
            "diploma_verification_code": "毕业证学信网在线验证码",
            "degree_number": "学位证编号",
            "degree_verification_code": "学位证学信网在线验证码"
        }}
    ],
    "work_experience": [
        {{
            "start_date": "工作开始日期（YYYY-MM-DD）",
            "end_date": "工作结束日期（YYYY-MM-DD，如果是当前工作则填null）",
            "company": "单位名称",
            "position": "岗位/职务",
            "is_psbc_independent_dev": "是否邮储银行自主研发工作经验（true/false）"
        }}
    ],
    "project_experience": [
        {{
            "start_date": "项目开始日期（YYYY-MM-DD）",
            "end_date": "项目结束日期（YYYY-MM-DD，如果是当前项目则填null）",
            "project_name": "项目名称",
            "description": "项目描述",
            "role": "担任角色/职责",
            "is_psbc_independent_dev": "是否邮储银行自主研发工作经验（true/false）"
        }}
    ],
    "technical_skills": {{
        "programming_languages": ["掌握的编程语言列表"],
        "skills": ["掌握的技能列表"],
        "certifications": ["专业证书列表"]
    }}
}}

请只返回 JSON，不要包含其他解释性文字。确保 JSON 格式有效。"""

    max_retries = 3
    for attempt in range(max_retries):
        try:
            response = completion(
                model=AI_MODEL,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                timeout=AI_TIMEOUT,
                api_base=DEEPSEEK_BASE_URL,
            )
            
            result_text = response.choices[0].message.content.strip()
            
            # 尝试提取 JSON（处理可能的 markdown 代码块包装）
            if result_text.startswith("```"):
                # 移除 markdown 代码块标记
                lines = result_text.split("\n")
                start_idx = 1 if lines[0].startswith("```") else 0
                end_idx = -1 if lines[-1].strip() == "```" else len(lines)
                result_text = "\n".join(lines[start_idx:end_idx])
            
            # 解析 JSON
            parsed_result = json.loads(result_text)
            return parsed_result
            
        except json.JSONDecodeError as e:
            print(f"JSON 解析错误: {e}")
            print(f"原始响应: {result_text}")
            return {"error": "JSON 解析失败", "raw_response": result_text}
        except Exception as e:
            print(f"API 调用错误 (尝试 {attempt + 1}/{max_retries}): {e}")
            if attempt < max_retries - 1:
                wait_time = (attempt + 1) * 3
                print(f"等待 {wait_time} 秒后重试...")
                time.sleep(wait_time)
            else:
                return {"error": str(e)}


def main():
    import sys
    
    # 支持命令行参数传入简历文件路径
    if len(sys.argv) > 1:
        file_path = sys.argv[1]
        # 如果是相对路径，转换为绝对路径
        if not os.path.isabs(file_path):
            file_path = os.path.join(os.path.dirname(__file__), file_path)
    else:
        # 默认路径（用于调试）
        file_path = "/Users/polly/Downloads/Sublime_Workspace/GitHub_Workspace/Hr_Copilot_v4/Resumes/王保利_20250706_191243.pdf"
    
    if not os.path.exists(file_path):
        print(f"错误: 文件不存在 - {file_path}")
        sys.exit(1)
    
    # 检查文件格式
    suffix = Path(file_path).suffix.lower()
    if suffix not in ['.pdf', '.docx']:
        print(f"错误: 不支持的文件格式 {suffix}，支持的格式: .pdf, .docx")
        sys.exit(1)
    
    print("=" * 60)
    print("简历解析开始")
    print("=" * 60)
    print(f"解析文件: {file_path}")
    print(f"文件格式: {suffix}")
    
    # 步骤1: 提取文本
    print(f"\n[1] 正在提取文本...")
    try:
        resume_text = extract_text_from_resume(file_path)
        print(f"提取完成，共 {len(resume_text)} 字符")
    except ValueError as e:
        print(f"错误: {e}")
        sys.exit(1)
    
    # 打印提取的文本（调试用）
    print("\n--- 提取的文本内容 ---")
    print(resume_text[:2000] + "..." if len(resume_text) > 2000 else resume_text)
    print("--- 文本内容结束 ---\n")
    
    # 步骤2: 使用 LLM 解析
    print("\n[2] 正在调用 DeepSeek 解析简历...")
    result = parse_resume_with_llm(resume_text)
    
    # 步骤3: 输出结果
    print("\n[3] 解析结果：")
    print("=" * 60)
    print(json.dumps(result, ensure_ascii=False, indent=2))
    print("=" * 60)
    
    # 保存结果到文件
    output_path = Path(file_path).parent / f"{Path(file_path).stem}_parsed.json"
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(result, f, ensure_ascii=False, indent=2)
    print(f"\n结果已保存到: {output_path}")
    
    return result


if __name__ == "__main__":
    main()
